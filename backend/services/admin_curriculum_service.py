"""Service layer for admin curriculum ingestion and governance workflows."""

from __future__ import annotations

import json
import os
import re
import shutil
import tempfile
from dataclasses import dataclass
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path
from uuid import UUID, uuid5

from sqlalchemy.orm import Session

from backend.core.config import settings
from backend.repositories.admin_curriculum_repo import AdminCurriculumRepository
from backend.repositories.neo4j_graph_repo import (
    Neo4jGraphConfig,
    Neo4jGraphRepository,
    Neo4jGraphRepositoryError,
)
from backend.schemas.admin_curriculum_schema import (
    ConceptInspectResponse,
    ConceptInspectTopicOut,
    CurriculumIngestionStatusResponse,
    CurriculumBulkIngestRequest,
    CurriculumBulkIngestResponse,
    CurriculumBulkScopeResult,
    CurriculumUploadRequest,
    CurriculumUploadResponse,
    CurriculumVersionActionRequest,
    CurriculumVersionActionResponse,
    PendingApprovalsResponse,
    TopicConceptMapOut,
    TopicInspectResponse,
    TopicMapPatchRequest,
)
from backend.services.rag_retrieve_service import (
    QdrantRuntimeConfig,
    QdrantVectorStore,
    RagRetrieveServiceError,
)


class AdminCurriculumServiceError(RuntimeError):
    pass


class AdminCurriculumValidationError(ValueError):
    pass


class AdminCurriculumNotFoundError(LookupError):
    pass


@dataclass(frozen=True)
class ConceptSection:
    concept_id: str
    concept_label: str
    prereq_concept_ids: list[str]
    confidence: float
    chunks: list[str]


@dataclass(frozen=True)
class ChunkedDocument:
    source_id: str
    topic_id: UUID
    topic_title: str
    concept_sections: list[ConceptSection]


class AdminCurriculumService:
    CHUNK_WORD_SIZE = 190
    CHUNK_OVERLAP_WORDS = 40
    MAX_CONCEPTS_PER_DOC = 12
    _CHUNK_NAMESPACE = UUID("f5cfde7f-0b62-46bf-bf17-838311a90be4")

    def __init__(self, db: Session):
        self.db = db
        self.repo = AdminCurriculumRepository(db)
        self.vector_store = QdrantVectorStore(
            QdrantRuntimeConfig(
                url=settings.qdrant_url,
                api_key=settings.qdrant_api_key or None,
                collection=settings.qdrant_collection,
                embedding_model=settings.qdrant_embedding_model,
            )
        )

    @staticmethod
    def _normalize_text(value: str) -> str:
        return re.sub(r"\s+", " ", value or "").strip().lower()

    @staticmethod
    def _is_truthy_env(value: str | None) -> bool:
        if value is None:
            return False
        return value.strip().lower() in {"1", "true", "yes", "on"}

    @classmethod
    def _is_heading(cls, line: str) -> bool:
        compact = line.strip()
        if len(compact) < 4 or len(compact) > 100:
            return False
        if compact.isupper() and len(compact.split()) <= 14:
            return True
        return bool(re.match(r"^(chapter|topic|unit|week)\b", compact, flags=re.IGNORECASE))

    @classmethod
    def _chunk_text(cls, text: str) -> list[str]:
        words = text.split()
        if not words:
            return []
        if len(words) <= cls.CHUNK_WORD_SIZE:
            return [" ".join(words)]

        chunks: list[str] = []
        start = 0
        while start < len(words):
            end = min(start + cls.CHUNK_WORD_SIZE, len(words))
            chunks.append(" ".join(words[start:end]))
            if end >= len(words):
                break
            start = max(end - cls.CHUNK_OVERLAP_WORDS, start + 1)
        return chunks

    @classmethod
    def _split_sections(cls, text: str) -> list[tuple[str, str]]:
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not lines:
            return []

        sections: list[tuple[str, str]] = []
        current_heading = ""
        buffer: list[str] = []
        for line in lines:
            if cls._is_heading(line) and buffer:
                sections.append((current_heading, "\n".join(buffer)))
                current_heading = line
                buffer = [line]
            else:
                if cls._is_heading(line):
                    current_heading = line
                buffer.append(line)
        if buffer:
            sections.append((current_heading, "\n".join(buffer)))
        return sections

    @staticmethod
    def _slugify(value: str, *, max_length: int = 48) -> str:
        slug = re.sub(r"[^a-z0-9]+", "-", (value or "").lower()).strip("-")
        slug = re.sub(r"-{2,}", "-", slug)
        if not slug:
            slug = "concept"
        return slug[:max_length].strip("-") or "concept"

    @classmethod
    def _section_label(cls, *, heading: str, section_text: str, topic_title: str, index: int) -> str:
        clean_heading = re.sub(r"^[\d\W_]+", "", (heading or "")).strip(" -:\t")
        if clean_heading:
            return clean_heading

        sentence = re.split(r"[.!?]", section_text)[0].strip()
        if sentence:
            words = sentence.split()
            trimmed = " ".join(words[:8]).strip()
            if len(trimmed) >= 4:
                return trimmed
        return f"{topic_title} concept {index + 1}"

    @classmethod
    def _scoped_concept_ids(
        cls,
        *,
        labels: list[str],
        subject: str,
        sss_level: str,
        term: int,
    ) -> list[str]:
        ids: list[str] = []
        seen: dict[str, int] = {}
        scope_prefix = f"{subject}:{sss_level.lower()}:t{term}"
        for label in labels:
            slug = cls._slugify(label)
            base = f"{scope_prefix}:{slug}"[:120]
            sequence = seen.get(base, 0) + 1
            seen[base] = sequence
            concept_id = base if sequence == 1 else f"{base}-{sequence}"
            ids.append(concept_id[:128])
        return ids

    @staticmethod
    def _extract_json_object(raw: str) -> dict | None:
        content = (raw or "").strip()
        if not content:
            return None
        try:
            parsed = json.loads(content)
            return parsed if isinstance(parsed, dict) else None
        except json.JSONDecodeError:
            pass
        match = re.search(r"\{[\s\S]*\}", content)
        if not match:
            return None
        try:
            parsed = json.loads(match.group(0))
        except json.JSONDecodeError:
            return None
        return parsed if isinstance(parsed, dict) else None

    def _maybe_refine_concept_labels_with_llm(
        self,
        *,
        topic_title: str,
        subject: str,
        sss_level: str,
        term: int,
        labels: list[str],
    ) -> list[str]:
        if not labels:
            return labels
        if not self._is_truthy_env(os.getenv("CURRICULUM_CONCEPT_USE_LLM")):
            return labels

        provider = os.getenv("LLM_PROVIDER", "groq").strip().lower()
        model = os.getenv("CURRICULUM_CONCEPT_LLM_MODEL", os.getenv("LLM_MODEL", "openai/gpt-oss-20b")).strip()

        api_key = (
            os.getenv("GROQ_API_KEY")
            if provider == "groq"
            else (os.getenv("LLM_API_KEY") or os.getenv("OPENAI_API_KEY"))
        )
        if not api_key:
            return labels

        try:
            from openai import OpenAI
        except ModuleNotFoundError:
            return labels

        base_url = ""
        if provider == "groq":
            base_url = os.getenv("GROQ_BASE_URL", "https://api.groq.com/openai/v1").strip()
        elif provider not in {"openai", ""} and os.getenv("LLM_API_BASE"):
            base_url = os.getenv("LLM_API_BASE", "").strip()

        try:
            client = OpenAI(api_key=api_key, base_url=base_url or None)
            prompt = (
                "You are normalizing curriculum concept labels.\n"
                f"Scope: {subject} {sss_level} term {term}. Topic: {topic_title}.\n"
                "Given candidate labels, return compact pedagogical concept names in JSON only.\n"
                "Keep meaning unchanged. Do not add or remove items.\n"
                'Return format: {"concepts":["..."]}\n'
                f"Candidates: {json.dumps(labels, ensure_ascii=True)}"
            )
            response = client.chat.completions.create(
                model=model,
                temperature=0,
                messages=[{"role": "user", "content": prompt}],
            )
            content = ""
            if response.choices:
                content = str(response.choices[0].message.content or "")
            parsed = self._extract_json_object(content)
            if not parsed:
                return labels
            concepts = parsed.get("concepts")
            if not isinstance(concepts, list) or len(concepts) != len(labels):
                return labels
            cleaned = [re.sub(r"\s+", " ", str(item)).strip() for item in concepts]
            if not all(cleaned):
                return labels
            return cleaned
        except Exception:
            return labels

    def _extract_concept_labels_with_llm(
        self,
        *,
        topic_title: str,
        subject: str,
        sss_level: str,
        term: int,
        raw_text: str,
    ) -> list[str] | None:
        should_use_llm = self._is_truthy_env(os.getenv("CURRICULUM_CONCEPT_EXTRACT_USE_LLM")) or self._is_truthy_env(
            os.getenv("CURRICULUM_CONCEPT_USE_LLM")
        )
        if not should_use_llm:
            return None

        provider = os.getenv("LLM_PROVIDER", "groq").strip().lower()
        model = os.getenv(
            "CURRICULUM_CONCEPT_EXTRACT_MODEL",
            os.getenv("CURRICULUM_CONCEPT_LLM_MODEL", os.getenv("LLM_MODEL", "openai/gpt-oss-20b")),
        ).strip()

        api_key = (
            os.getenv("GROQ_API_KEY")
            if provider == "groq"
            else (os.getenv("LLM_API_KEY") or os.getenv("OPENAI_API_KEY"))
        )
        if not api_key:
            return None

        try:
            from openai import OpenAI
        except ModuleNotFoundError:
            return None

        base_url = ""
        if provider == "groq":
            base_url = os.getenv("GROQ_BASE_URL", "https://api.groq.com/openai/v1").strip()
        elif provider not in {"openai", ""} and os.getenv("LLM_API_BASE"):
            base_url = os.getenv("LLM_API_BASE", "").strip()

        text_sample = re.sub(r"\s+", " ", raw_text).strip()[:12000]
        if not text_sample:
            return None

        try:
            client = OpenAI(api_key=api_key, base_url=base_url or None)
            prompt = (
                "You are extracting curriculum concepts from study notes.\n"
                f"Scope: subject={subject}, level={sss_level}, term={term}, topic={topic_title}.\n"
                "Return JSON only in this exact format: {\"concepts\": [\"...\"]}.\n"
                "Extract 3 to 10 concise concept names relevant to this topic.\n"
                "Do not include explanations, numbering, or markdown.\n\n"
                f"Notes:\n{text_sample}"
            )
            response = client.chat.completions.create(
                model=model,
                temperature=0,
                messages=[{"role": "user", "content": prompt}],
            )
            content = ""
            if response.choices:
                content = str(response.choices[0].message.content or "")
            parsed = self._extract_json_object(content)
            if not parsed:
                return None
            concepts = parsed.get("concepts")
            if not isinstance(concepts, list):
                return None

            cleaned: list[str] = []
            seen: set[str] = set()
            for item in concepts:
                normalized = re.sub(r"\s+", " ", str(item)).strip()
                if not normalized:
                    continue
                key = normalized.lower()
                if key in seen:
                    continue
                seen.add(key)
                cleaned.append(normalized)

            if not (3 <= len(cleaned) <= 10):
                return None
            return cleaned[: self.MAX_CONCEPTS_PER_DOC]
        except Exception:
            return None

    @staticmethod
    def _read_docx(path: Path) -> str:
        try:
            from docx import Document
        except ModuleNotFoundError as exc:
            raise AdminCurriculumServiceError(
                "python-docx is not installed. Add `python-docx` to backend dependencies."
            ) from exc

        doc = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text and paragraph.text.strip()]
        return "\n".join(paragraphs)

    @staticmethod
    def _read_text_file(path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="ignore")

    @staticmethod
    def _collect_supported_files(root: Path) -> tuple[list[Path], list[Path]]:
        all_files = [path for path in root.rglob("*") if path.is_file()]
        supported = [path for path in all_files if path.suffix.lower() in {".docx", ".txt"}]
        skipped = [path for path in all_files if path.suffix.lower() not in {".docx", ".txt"}]
        return supported, skipped

    @staticmethod
    def _topic_hint_from_file_name(file_path: Path) -> str:
        stem = file_path.stem.upper()
        tokens = re.split(r"[-_\s]+", stem)
        skip_tokens = {
            "COMPLETE",
            "FIRST",
            "SECOND",
            "THIRD",
            "TERM",
            "SS1",
            "SS2",
            "SS3",
            "SSS1",
            "SSS2",
            "SSS3",
            "MATHEMATICS",
            "MATH",
            "ENGLISH",
            "CIVIC",
            "EDUCATION",
            "NOTES",
            "NOTE",
        }
        filtered = [token for token in tokens if token and token not in skip_tokens]
        if not filtered:
            return stem
        return " ".join(filtered)

    @classmethod
    def _best_topic_match(cls, *, topic_hint: str, topics: list) -> tuple[UUID, str] | None:
        if not topics:
            return None
        normalized_hint = cls._normalize_text(topic_hint)
        if not normalized_hint:
            return None

        best_topic = None
        best_score = -1.0
        for topic in topics:
            title = str(topic.title)
            normalized_title = cls._normalize_text(title)
            score = SequenceMatcher(a=normalized_hint, b=normalized_title).ratio()
            if normalized_hint in normalized_title or normalized_title in normalized_hint:
                score = max(score, 0.92)
            if score > best_score:
                best_score = score
                best_topic = topic

        if best_topic is None or best_score < 0.35:
            return None
        return best_topic.id, str(best_topic.title)

    @classmethod
    def _version_name_default(cls, *, subject: str, sss_level: str, term: int) -> str:
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
        return f"{subject}-{sss_level}-term{term}-{timestamp}"

    @staticmethod
    def _infer_subject(scope_text: str) -> str | None:
        value = scope_text.upper()
        if re.search(r"\b(MATHEMATICS|MATHS|MATH)\b", value):
            return "math"
        if re.search(r"\bENGLISH\b", value):
            return "english"
        if re.search(r"\bCIVIC\b", value):
            return "civic"
        return None

    @staticmethod
    def _infer_sss_level(scope_text: str) -> str | None:
        value = scope_text.upper()
        patterns = [
            r"\bSSS?\s*([123])\b",  # SSS1, SS1, SS 1
            r"\bS([123])\b",        # S1, S2, S3
        ]
        for pattern in patterns:
            match = re.search(pattern, value)
            if match:
                return f"SSS{match.group(1)}"
        return None

    @staticmethod
    def _infer_term(scope_text: str) -> int | None:
        value = scope_text.upper()
        if (
            "FIRST TERM" in value
            or "1ST TERM" in value
            or "TERM 1" in value
            or "TERM ONE" in value
        ):
            return 1
        if (
            "SECOND TERM" in value
            or "2ND TERM" in value
            or "TERM 2" in value
            or "TERM TWO" in value
        ):
            return 2
        if (
            "THIRD TERM" in value
            or "3RD TERM" in value
            or "TERM 3" in value
            or "TERM THREE" in value
        ):
            return 3
        return None

    @classmethod
    def _infer_scope_from_file(cls, *, root: Path, file_path: Path) -> tuple[str, str, int] | None:
        try:
            relative_text = str(file_path.relative_to(root))
        except ValueError:
            relative_text = str(file_path)
        candidate = f"{relative_text} {file_path.stem}"
        subject = cls._infer_subject(candidate)
        sss_level = cls._infer_sss_level(candidate)
        term = cls._infer_term(candidate)
        if not subject or not sss_level or not term:
            return None
        return (subject, sss_level, term)

    @classmethod
    def _discover_scoped_file_groups(
        cls,
        *,
        source_root: Path,
    ) -> tuple[dict[tuple[str, str, int], list[Path]], list[str], list[str]]:
        supported, skipped = cls._collect_supported_files(source_root)
        grouped: dict[tuple[str, str, int], list[Path]] = {}
        undetected: list[str] = []
        for file_path in supported:
            scope = cls._infer_scope_from_file(root=source_root, file_path=file_path)
            if scope is None:
                undetected.append(str(file_path.relative_to(source_root)))
                continue
            grouped.setdefault(scope, []).append(file_path)
        skipped_rel = [str(path.relative_to(source_root)) for path in skipped]
        return grouped, skipped_rel, undetected

    def _extract_document_chunks(
        self,
        *,
        file_path: Path,
        scope_topics: list,
        subject: str,
        sss_level: str,
        term: int,
    ) -> ChunkedDocument | None:
        if file_path.suffix.lower() == ".docx":
            text = self._read_docx(file_path)
        elif file_path.suffix.lower() == ".txt":
            text = self._read_text_file(file_path)
        else:
            return None

        text = text.strip()
        if not text:
            return None

        topic_hint = self._topic_hint_from_file_name(file_path)
        match = self._best_topic_match(topic_hint=topic_hint, topics=scope_topics)
        if match is None:
            return None
        topic_id, topic_title = match

        section_pairs = self._split_sections(text)
        if not section_pairs:
            section_pairs = [("", text)]

        section_payloads: list[tuple[str, str, list[str], float]] = []
        for index, (heading, section_text) in enumerate(section_pairs):
            section_chunks = self._chunk_text(section_text)
            if not section_chunks:
                continue
            label = self._section_label(
                heading=heading,
                section_text=section_text,
                topic_title=topic_title,
                index=index,
            )
            confidence = 0.9 if heading else 0.72
            section_payloads.append((heading, label, section_chunks, confidence))

        if not section_payloads:
            return None

        # Keep ingestion bounded for very long notes.
        section_payloads = section_payloads[: self.MAX_CONCEPTS_PER_DOC]
        llm_labels = self._extract_concept_labels_with_llm(
            topic_title=topic_title,
            subject=subject,
            sss_level=sss_level,
            term=term,
            raw_text=text,
        )
        if llm_labels:
            target_count = min(len(section_payloads), len(llm_labels), self.MAX_CONCEPTS_PER_DOC)
            section_payloads = [
                (heading, llm_labels[idx], chunks, confidence)
                for idx, (heading, _, chunks, confidence) in enumerate(section_payloads[:target_count])
            ]
        else:
            labels = [item[1] for item in section_payloads]
            labels = self._maybe_refine_concept_labels_with_llm(
                topic_title=topic_title,
                subject=subject,
                sss_level=sss_level,
                term=term,
                labels=labels,
            )
            section_payloads = [
                (heading, labels[idx], chunks, confidence)
                for idx, (heading, _, chunks, confidence) in enumerate(section_payloads)
            ]

        if not section_payloads:
            return None

        labels = [item[1] for item in section_payloads]
        concept_ids = self._scoped_concept_ids(
            labels=labels,
            subject=subject,
            sss_level=sss_level,
            term=term,
        )

        concept_sections: list[ConceptSection] = []
        for index, (_, label, chunks, confidence) in enumerate(section_payloads):
            concept_id = concept_ids[index]
            prereqs = [concept_ids[index - 1]] if index > 0 else []
            concept_sections.append(
                ConceptSection(
                    concept_id=concept_id,
                    concept_label=label,
                    prereq_concept_ids=prereqs,
                    confidence=confidence,
                    chunks=chunks,
                )
            )

        return ChunkedDocument(
            source_id=file_path.name,
            topic_id=topic_id,
            topic_title=topic_title,
            concept_sections=concept_sections,
        )

    def upload_curriculum(
        self,
        *,
        payload: CurriculumUploadRequest,
        actor_user_id: UUID | None = None,
    ) -> CurriculumUploadResponse:
        source_root = Path(payload.source_root).expanduser().resolve()
        if not source_root.exists() or not source_root.is_dir():
            raise AdminCurriculumValidationError(f"source_root does not exist or is not a directory: {source_root}")

        scope_topics = self.repo.get_scope_topics(
            subject=payload.subject,
            sss_level=payload.sss_level,
            term=payload.term,
        )
        if not scope_topics:
            raise AdminCurriculumValidationError(
                "No topics found for provided scope. Seed topics before curriculum ingestion."
            )

        supported_files, skipped_files = self._collect_supported_files(source_root)
        if not supported_files:
            raise AdminCurriculumValidationError(
                "No supported files found in source_root. Supported extensions: .docx, .txt"
            )

        version_name = payload.version_name or self._version_name_default(
            subject=payload.subject,
            sss_level=payload.sss_level,
            term=payload.term,
        )
        if self.repo.get_curriculum_version_by_name(version_name):
            raise AdminCurriculumValidationError(f"Curriculum version name already exists: {version_name}")

        version = self.repo.create_curriculum_version(
            version_name=version_name,
            subject=payload.subject,
            sss_level=payload.sss_level,
            term=payload.term,
            source_root=str(source_root),
            uploaded_by=actor_user_id,
            status="ingesting",
            source_file_count=len(supported_files),
            metadata_payload={"skipped_extensions": sorted({path.suffix.lower() for path in skipped_files})},
        )
        job = self.repo.create_ingestion_job(
            version_id=version.id,
            created_by=actor_user_id,
            status="parsing",
            current_stage="parsing",
        )
        self.repo.append_ingestion_log(
            job,
            stage="parsing",
            message="Started ingestion",
            extra={"source_root": str(source_root), "supported_files": len(supported_files)},
        )
        self.db.commit()

        try:
            chunk_rows: list[dict] = []
            mapped_topic_ids: set[UUID] = set()
            mapped_concept_ids: set[str] = set()
            topics_for_neo4j: dict[str, dict] = {}
            topic_to_concepts: dict[str, dict] = {}
            all_concepts_for_neo4j: dict[str, str] = {}
            prereq_edges: set[tuple[str, str]] = set()
            processed_file_count = 0
            processed_chunks = 0

            for index, file_path in enumerate(supported_files, start=1):
                parsed = self._extract_document_chunks(
                    file_path=file_path,
                    scope_topics=scope_topics,
                    subject=payload.subject,
                    sss_level=payload.sss_level,
                    term=payload.term,
                )
                if parsed is None:
                    self.repo.append_ingestion_log(
                        job,
                        stage="parsing",
                        message="Skipped file (empty/unmatched topic)",
                        extra={"file": file_path.name},
                    )
                    continue

                processed_file_count += 1
                mapped_topic_ids.add(parsed.topic_id)
                topic_id_str = str(parsed.topic_id)
                topics_for_neo4j.setdefault(
                    topic_id_str,
                    {
                        "topic_id": topic_id_str,
                        "title": parsed.topic_title,
                        "sss_level": payload.sss_level,
                        "term": payload.term,
                    },
                )
                topic_bundle = topic_to_concepts.setdefault(
                    topic_id_str,
                    {"title": parsed.topic_title, "concept_ids": [], "concept_labels": {}},
                )
                for concept in parsed.concept_sections:
                    mapped_concept_ids.add(concept.concept_id)
                    if concept.concept_id not in topic_bundle["concept_ids"]:
                        topic_bundle["concept_ids"].append(concept.concept_id)
                    topic_bundle["concept_labels"][concept.concept_id] = concept.concept_label
                    all_concepts_for_neo4j[concept.concept_id] = concept.concept_label
                    for prereq in concept.prereq_concept_ids:
                        prereq_id = str(prereq).strip()
                        if prereq_id and prereq_id != concept.concept_id:
                            prereq_edges.add((prereq_id, concept.concept_id))
                    self.repo.upsert_topic_map(
                        version_id=version.id,
                        topic_id=parsed.topic_id,
                        concept_id=concept.concept_id,
                        prereq_concept_ids=list(concept.prereq_concept_ids),
                        confidence=float(concept.confidence),
                        is_manual_override=False,
                        created_by=actor_user_id,
                    )

                    for chunk_index, chunk_text in enumerate(concept.chunks):
                        deterministic_id = uuid5(
                            self._CHUNK_NAMESPACE,
                            f"{version.id}:{parsed.source_id}:{parsed.topic_id}:{concept.concept_id}:{chunk_index}",
                        )
                        chunk_payload = {
                            "chunk_id": str(deterministic_id),
                            "source_id": parsed.source_id,
                            "text": chunk_text,
                            "subject": payload.subject,
                            "sss_level": payload.sss_level,
                            "term": payload.term,
                            "topic_id": str(parsed.topic_id),
                            "topic_title": parsed.topic_title,
                            "concept_id": concept.concept_id,
                            "concept_label": concept.concept_label,
                            "curriculum_version_id": str(version.id),
                            "approved": False,
                            "chunk_index": chunk_index,
                            "citation_topic_title": parsed.topic_title,
                            "citation_source_id": parsed.source_id,
                            "citation_chunk_index": chunk_index,
                            "citation_concept_label": concept.concept_label,
                        }
                        chunk_rows.append({"id": deterministic_id, "text": chunk_text, "payload": chunk_payload})

                    processed_chunks += len(concept.chunks)
                progress = int((index / max(len(supported_files), 1)) * 55)
                self.repo.update_ingestion_job(
                    job,
                    progress_percent=min(progress, 55),
                    processed_files_count=processed_file_count,
                    processed_chunks_count=processed_chunks,
                )

            self.repo.update_ingestion_job(
                job,
                status="embedding",
                current_stage="embedding",
                progress_percent=60,
                processed_files_count=processed_file_count,
                processed_chunks_count=processed_chunks,
            )
            self.repo.append_ingestion_log(
                job,
                stage="embedding",
                message="Embedding and indexing chunks",
                extra={"chunk_count": len(chunk_rows)},
            )

            if chunk_rows:
                self.vector_store.upsert_chunks(chunk_rows)

            should_sync_neo4j = settings.use_neo4j_graph and bool(
                settings.neo4j_uri and settings.neo4j_user and settings.neo4j_password
            )
            if should_sync_neo4j and topic_to_concepts:
                neo_repo = Neo4jGraphRepository(
                    Neo4jGraphConfig(
                        uri=settings.neo4j_uri,
                        user=settings.neo4j_user,
                        password=settings.neo4j_password,
                    )
                )
                try:
                    neo_repo.ensure_subject_topics(subject=payload.subject, topics=list(topics_for_neo4j.values()))
                    flat_concept_ids: list[str] = []
                    for topic_id, bundle in topic_to_concepts.items():
                        concept_ids = [str(concept_id) for concept_id in bundle["concept_ids"]]
                        flat_concept_ids.extend(concept_ids)
                        neo_repo.ensure_topic_concept_links(
                            subject=payload.subject,
                            sss_level=payload.sss_level,
                            term=payload.term,
                            topic_id=topic_id,
                            topic_title=str(bundle["title"]),
                            concept_ids=concept_ids,
                            concept_labels=dict(bundle["concept_labels"]),
                        )
                    neo_repo.ensure_concepts_with_labels(
                        subject=payload.subject,
                        sss_level=payload.sss_level,
                        term=payload.term,
                        concepts=[
                            {"id": concept_id, "name": concept_label}
                            for concept_id, concept_label in all_concepts_for_neo4j.items()
                        ],
                    )
                    if prereq_edges:
                        neo_repo.ensure_prerequisite_edges(edges=sorted(prereq_edges))
                    else:
                        neo_repo.ensure_prerequisite_chain(concept_ids=flat_concept_ids)
                except Neo4jGraphRepositoryError as exc:
                    raise AdminCurriculumServiceError(f"Neo4j ingestion sync failed: {exc}") from exc
                finally:
                    neo_repo.close()

            topic_ids_list = list(mapped_topic_ids)
            affected_topics = self.repo.set_topics_version(
                topic_ids=topic_ids_list,
                version_id=version.id,
                is_approved=False,
            )

            version_metadata = dict(version.metadata_payload or {})
            version_metadata.update(
                {
                    "processed_files_count": processed_file_count,
                    "processed_chunks_count": processed_chunks,
                    "processed_concepts_count": len(mapped_concept_ids),
                    "indexed_at": datetime.now(timezone.utc).isoformat(),
                    "affected_topics": affected_topics,
                }
            )

            self.repo.update_curriculum_version(
                version,
                status="pending_approval" if chunk_rows else "failed",
                metadata_payload=version_metadata,
            )
            self.repo.update_ingestion_job(
                job,
                status="completed" if chunk_rows else "failed",
                current_stage="completed" if chunk_rows else "failed",
                progress_percent=100,
                processed_files_count=processed_file_count,
                processed_chunks_count=processed_chunks,
                error_message=None if chunk_rows else "No indexable chunks produced",
                finished_at=datetime.now(timezone.utc),
            )
            self.db.commit()

            return CurriculumUploadResponse(
                version_id=version.id,
                job_id=job.id,
                status=version.status,
                discovered_files=len(supported_files),
                skipped_files=len(skipped_files),
                processed_chunks=processed_chunks,
            )
        except Exception as exc:
            self.db.rollback()
            persisted_job = self.repo.get_ingestion_job(job.id)
            persisted_version = self.repo.get_curriculum_version(version.id)
            if persisted_job and persisted_version:
                self.repo.update_ingestion_job(
                    persisted_job,
                    status="failed",
                    current_stage="failed",
                    progress_percent=min(persisted_job.progress_percent, 99),
                    error_message=str(exc)[:2000],
                    finished_at=datetime.now(timezone.utc),
                )
                self.repo.update_curriculum_version(persisted_version, status="failed")
                self.db.commit()
            raise AdminCurriculumServiceError(f"Curriculum ingestion failed: {exc}") from exc

    def ingest_all_from_source_root(
        self,
        *,
        payload: CurriculumBulkIngestRequest,
        actor_user_id: UUID | None = None,
    ) -> CurriculumBulkIngestResponse:
        source_root = Path(payload.source_root).expanduser().resolve()
        if not source_root.exists() or not source_root.is_dir():
            raise AdminCurriculumValidationError(
                f"source_root does not exist or is not a directory: {source_root}"
            )

        grouped, skipped_files, undetected_scope_files = self._discover_scoped_file_groups(
            source_root=source_root
        )

        results: list[CurriculumBulkScopeResult] = []
        approve_ready_version_ids: list[UUID] = []

        for scope, files in sorted(grouped.items()):
            subject, sss_level, term = scope
            with tempfile.TemporaryDirectory(prefix="curriculum_scope_") as tmp_dir:
                temp_root = Path(tmp_dir)
                used_names: set[str] = set()
                for file_path in files:
                    candidate = file_path.name
                    if candidate in used_names:
                        base = file_path.stem
                        suffix = file_path.suffix
                        serial = 2
                        while f"{base}__{serial}{suffix}" in used_names:
                            serial += 1
                        candidate = f"{base}__{serial}{suffix}"
                    used_names.add(candidate)
                    shutil.copy2(file_path, temp_root / candidate)

                try:
                    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S%f")
                    response = self.upload_curriculum(
                        payload=CurriculumUploadRequest(
                            subject=subject,
                            sss_level=sss_level,
                            term=term,
                            source_root=str(temp_root),
                            version_name=f"auto-{subject}-{sss_level}-term{term}-{timestamp}",
                        ),
                        actor_user_id=actor_user_id,
                    )
                    results.append(
                        CurriculumBulkScopeResult(
                            subject=subject,  # type: ignore[arg-type]
                            sss_level=sss_level,  # type: ignore[arg-type]
                            term=term,
                            files_count=len(files),
                            version_id=response.version_id,
                            job_id=response.job_id,
                            status=response.status,
                        )
                    )
                    if response.status == "pending_approval":
                        approve_ready_version_ids.append(response.version_id)
                except Exception as exc:
                    results.append(
                        CurriculumBulkScopeResult(
                            subject=subject,  # type: ignore[arg-type]
                            sss_level=sss_level,  # type: ignore[arg-type]
                            term=term,
                            files_count=len(files),
                            status="failed",
                            message=str(exc),
                        )
                    )

        failed_scopes = len([item for item in results if item.status == "failed"])
        succeeded_scopes = len(results) - failed_scopes
        return CurriculumBulkIngestResponse(
            source_root=str(source_root),
            total_supported_files=sum(len(files) for files in grouped.values()) + len(undetected_scope_files),
            total_unsupported_files=len(skipped_files),
            total_undetected_scope_files=len(undetected_scope_files),
            discovered_scopes=len(grouped),
            succeeded_scopes=succeeded_scopes,
            failed_scopes=failed_scopes,
            approve_ready_version_ids=approve_ready_version_ids,
            skipped_files=sorted(skipped_files),
            undetected_scope_files=sorted(undetected_scope_files),
            results=results,
        )

    def get_ingestion_status(self, *, job_id: UUID | None = None) -> CurriculumIngestionStatusResponse:
        jobs = self.repo.list_ingestion_jobs(job_id=job_id)
        return CurriculumIngestionStatusResponse(
            jobs=[
                {
                    "id": job.id,
                    "version_id": job.version_id,
                    "status": job.status,
                    "progress_percent": job.progress_percent,
                    "current_stage": job.current_stage,
                    "processed_files_count": job.processed_files_count,
                    "processed_chunks_count": job.processed_chunks_count,
                    "error_message": job.error_message,
                    "created_at": job.created_at,
                    "updated_at": job.updated_at,
                }
                for job in jobs
            ]
        )

    def get_pending_approvals(self) -> PendingApprovalsResponse:
        versions = self.repo.list_pending_approvals()
        return PendingApprovalsResponse(
            versions=[
                {
                    "id": row.id,
                    "version_name": row.version_name,
                    "subject": row.subject,
                    "sss_level": row.sss_level,
                    "term": row.term,
                    "source_file_count": row.source_file_count,
                    "status": row.status,
                    "created_at": row.created_at,
                    "updated_at": row.updated_at,
                }
                for row in versions
            ]
        )

    def inspect_topic(self, *, topic_id: UUID) -> TopicInspectResponse:
        row = self.repo.get_topic_with_subject(topic_id)
        if row is None:
            raise AdminCurriculumNotFoundError(f"Topic not found: {topic_id}")
        topic, subject_slug = row
        maps = self.repo.get_topic_maps(topic_id=topic_id)
        return TopicInspectResponse(
            topic_id=topic.id,
            subject=subject_slug,  # type: ignore[arg-type]
            sss_level=topic.sss_level,  # type: ignore[arg-type]
            term=topic.term,
            title=topic.title,
            is_approved=topic.is_approved,
            curriculum_version_id=topic.curriculum_version_id,
            mappings=[
                TopicConceptMapOut(
                    id=item.id,
                    version_id=item.version_id,
                    topic_id=item.topic_id,
                    concept_id=item.concept_id,
                    prereq_concept_ids=list(item.prereq_concept_ids or []),
                    confidence=float(item.confidence),
                    is_manual_override=item.is_manual_override,
                    created_at=item.created_at,
                    updated_at=item.updated_at,
                )
                for item in maps
            ],
        )

    def inspect_concept(self, *, concept_id: str) -> ConceptInspectResponse:
        rows = self.repo.get_concept_maps(concept_id=concept_id)
        if not rows:
            raise AdminCurriculumNotFoundError(f"Concept not found in curriculum mappings: {concept_id}")

        prereq_ids: set[str] = set()
        topics: list[ConceptInspectTopicOut] = []
        for map_row, topic_row, subject_slug in rows:
            for prereq in (map_row.prereq_concept_ids or []):
                prereq_ids.add(str(prereq))
            topics.append(
                ConceptInspectTopicOut(
                    topic_id=topic_row.id,
                    title=topic_row.title,
                    subject=subject_slug,  # type: ignore[arg-type]
                    sss_level=topic_row.sss_level,  # type: ignore[arg-type]
                    term=topic_row.term,
                    confidence=float(map_row.confidence),
                )
            )
        return ConceptInspectResponse(
            concept_id=concept_id,
            prereq_concept_ids=sorted(prereq_ids),
            topics=topics,
        )

    def patch_topic_map(
        self,
        *,
        topic_id: UUID,
        payload: TopicMapPatchRequest,
        actor_user_id: UUID | None = None,
    ) -> TopicInspectResponse:
        topic_row = self.repo.get_topic_with_subject(topic_id)
        if topic_row is None:
            raise AdminCurriculumNotFoundError(f"Topic not found: {topic_id}")
        if self.repo.get_curriculum_version(payload.version_id) is None:
            raise AdminCurriculumNotFoundError(f"Curriculum version not found: {payload.version_id}")

        for item in payload.mappings:
            self.repo.upsert_topic_map(
                version_id=payload.version_id,
                topic_id=topic_id,
                concept_id=item.concept_id,
                prereq_concept_ids=item.prereq_concept_ids,
                confidence=item.confidence,
                is_manual_override=item.is_manual_override,
                created_by=actor_user_id,
            )
        self.db.commit()
        return self.inspect_topic(topic_id=topic_id)

    def approve_version(
        self,
        *,
        version_id: UUID,
        payload: CurriculumVersionActionRequest,
    ) -> CurriculumVersionActionResponse:
        version = self.repo.get_curriculum_version(version_id)
        if version is None:
            raise AdminCurriculumNotFoundError(f"Curriculum version not found: {version_id}")
        if version.status not in {"pending_approval", "approved", "published"}:
            raise AdminCurriculumValidationError(
                f"Version status must be pending_approval|approved|published before publish, got: {version.status}"
            )

        scope_topics = self.repo.get_scope_topics(
            subject=version.subject,
            sss_level=version.sss_level,
            term=version.term,
        )
        topic_ids = [topic.id for topic in scope_topics]
        affected_topics = self.repo.set_topics_version(topic_ids=topic_ids, version_id=version.id, is_approved=True)
        self.repo.update_curriculum_version(
            version,
            status="published",
            approved_by=payload.actor_user_id,
            approved_at=datetime.now(timezone.utc),
        )
        try:
            self.vector_store.set_approval_flag(curriculum_version_id=version.id, approved=True)
        except RagRetrieveServiceError:
            self.db.rollback()
            raise
        self.db.commit()
        return CurriculumVersionActionResponse(
            version_id=version.id,
            status=version.status,  # type: ignore[arg-type]
            affected_topics=affected_topics,
            message="Curriculum version published successfully.",
        )

    def rollback_version(
        self,
        *,
        version_id: UUID,
        payload: CurriculumVersionActionRequest,
    ) -> CurriculumVersionActionResponse:
        version = self.repo.get_curriculum_version(version_id)
        if version is None:
            raise AdminCurriculumNotFoundError(f"Curriculum version not found: {version_id}")

        scope_topics = self.repo.get_scope_topics(
            subject=version.subject,
            sss_level=version.sss_level,
            term=version.term,
        )
        topic_ids = [topic.id for topic in scope_topics]
        affected_topics = self.repo.set_topics_version(topic_ids=topic_ids, version_id=version.id, is_approved=False)
        self.repo.update_curriculum_version(version, status="rolled_back")

        previous = self.repo.get_latest_published_version_for_scope(
            subject=version.subject,
            sss_level=version.sss_level,
            term=version.term,
            exclude_version_id=version.id,
        )
        if previous is not None:
            self.repo.update_curriculum_version(previous, status="published")
            self.repo.set_topics_version(topic_ids=topic_ids, version_id=previous.id, is_approved=True)

        try:
            self.vector_store.set_approval_flag(curriculum_version_id=version.id, approved=False)
        except RagRetrieveServiceError:
            self.db.rollback()
            raise
        self.db.commit()

        return CurriculumVersionActionResponse(
            version_id=version.id,
            status=version.status,  # type: ignore[arg-type]
            affected_topics=affected_topics,
            message="Curriculum version rolled back.",
        )
